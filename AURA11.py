import streamlit as st
from docx import Document
import os



# List of valid Aura IDs (this can later be replaced with a real database query)
valid_aura_ids = ["aura_id_12345", "aura_id_67890", "aura_id_11223"]

# Function to calculate risk (simple example, you can adjust the logic as needed)
def calculate_risk(heart_rate, oxygen_saturation, blood_pressure):
    if heart_rate > 100 or oxygen_saturation < 90 or blood_pressure > '140/90':
        return "High Risk"
    elif heart_rate > 90 or oxygen_saturation < 95 or blood_pressure > '120/80':
        return "Medium Risk"
    else:
        return "Low Risk"

# Function to generate a Word report and return the file path
def generate_example_report(heart_rate, oxygen_saturation, blood_pressure, risk):
    # Create a new Document
    doc = Document()
    
    # Add a title
    doc.add_heading('Heart Health Report', 0)
    
    # Add heart rate, oxygen saturation, blood pressure, and risk
    doc.add_paragraph(f'Heart Rate: {heart_rate} bpm')
    doc.add_paragraph(f'Oxygen Saturation: {oxygen_saturation}%')
    doc.add_paragraph(f'Blood Pressure: {blood_pressure}')
    doc.add_paragraph(f'Risk: {risk}')
    
    # Save the document
    report_path = "example_heart_health_report.docx"
    doc.save(report_path)
    return report_path


# Streamlit interface
st.title("Heart Attack Prediction Dashboard")

# Step 1: User Login using Aura ID
aura_id = st.text_input("Enter your Aura ID:")

if aura_id in valid_aura_ids:
    st.success(f"Welcome, {aura_id}! You are logged in.")
    
    # Step 2: User Inputs for Heart Rate, Oxygen Saturation, and Blood Pressure
    heart_rate = st.number_input("Enter your Heart Rate (bpm):", min_value=0)
    oxygen_saturation = st.number_input("Enter your Oxygen Saturation (%):", min_value=0)
    blood_pressure = st.text_input("Enter your Blood Pressure (e.g., 120/80):")

    # Step 3: Calculate the risk based on inputs
    if heart_rate and oxygen_saturation and blood_pressure:
        risk = calculate_risk(heart_rate, oxygen_saturation, blood_pressure)

        # Display the entered data and risk on the website
        st.header("Heart Health Information")
        st.write(f"**Heart Rate**: {heart_rate} bpm")
        st.write(f"**Oxygen Saturation**: {oxygen_saturation}%")
        st.write(f"**Blood Pressure**: {blood_pressure}")
        st.write(f"**Risk**: {risk}")

        # Step 4: Generate the Word report and provide download option
        st.header("Download Your Heart Health Report")
        report_path = generate_example_report(heart_rate, oxygen_saturation, blood_pressure, risk)

        # Provide the download link for the Word file
        with open(report_path, "rb") as f:
            st.download_button(
                label="Download Report as Word File",
                data=f,
                file_name=report_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
else:
    st.warning("Invalid Aura ID. Please try again.")
